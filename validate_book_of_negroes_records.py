import pandas as pd
from docx import Document
import os
import re

# --- MAPPING CONFIGURATION ---
FILE_TO_BOOK_MAP = {
    "Book_One_Part_One_of_the_Book_of_Negroes.docx": "Book One Part One",
    "Book_One_Part_Two_of_the_Book_of_Negroes.docx": "Book One Part Two",
    "Book_Two.docx": "Book Two",
    "Book_Three.docx": "Book Three"
}

# Ordered list of documents for backtracking across files
ORDERED_FILES = [
    "Book_One_Part_One_of_the_Book_of_Negroes.docx",
    "Book_One_Part_Two_of_the_Book_of_Negroes.docx",
    "Book_Two.docx",
    "Book_Three.docx"
]

def extract_header_info(line):
    """
    Extracts Ship Name, City, and Commander from a ship header line.
    Returns a dict with 'ship', 'city', 'cmd' or None if not a ship line.
    """
    if "bound for" in line.lower():
        parts = re.split(r"\s+bound\s+for\s+", line, flags=re.I)
        ship_part = parts[0].strip()
        rest = parts[1].strip()

        prefix_pattern = r"^(Ship|Brig|Sloop|Schooner|Brigantine|Snow)\s+(.*)$"
        m_prefix = re.match(prefix_pattern, ship_part, re.I)
        ship_name = m_prefix.group(2).strip() if m_prefix else ship_part

        cities = [
            "River St. John's", "St. John's River", "St. John's",
            "Port Roseway", "Halifax", "Annapolis Royal",
            "Spithead & Germany", "Shelburne", "River St. Johns"
        ]

        found_city, commander = "-", "-"
        for city in sorted(cities, key=len, reverse=True):
            if re.search(re.escape(city), rest, re.I):
                found_city = city
                cmd_raw = re.split(re.escape(city), rest, flags=re.I)[-1].strip()
                if cmd_raw:
                    cmd_raw = re.sub(r",?\s*Master$", "", cmd_raw, flags=re.I).strip()
                    commander = cmd_raw
                break

        if found_city == "-":
            found_city = rest

        return {
            "ship": ship_name if ship_name else None,
            "city": found_city,
            "cmd": commander
        }

    # Check for "On Board the Ship... Master" pattern without "bound for"
    m_master = re.search(
        r"(?:On\s+Board\s+the\s+)?(?:Ship|Brig|Sloop|Schooner)\s+(.*?)\s+([A-Z][A-Za-z\.]+(?:\s+[A-Z][A-Za-z\.]+)*),\s+Master",
        line, re.I
    )
    if m_master:
        return {
            "ship": m_master.group(1).strip(),
            "cmd": m_master.group(2).strip(),
            "city": "-"
        }

    return None


def get_word_content_ordered(directory_path):
    """
    Scrapes text from Word documents IN ORDER defined by ORDERED_FILES.
    Returns a flat ordered list of records across all files for backtracking.
    """
    all_records = []

    for filename in ORDERED_FILES:
        file_path = os.path.join(directory_path, filename)
        if not os.path.exists(file_path):
            print(f"Warning: '{filename}' not found in '{directory_path}', skipping.")
            continue

        try:
            doc = Document(file_path)
            entries = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            entries.append(cell.text.strip())

            for entry in entries:
                all_records.append({
                    "Notes": entry,
                    "Source_Word_File": filename
                })
        except Exception as e:
            print(f"Error reading {filename}: {e}")

    return all_records


def find_last_header_before(index, all_records):
    """
    Backtracks from the given index through all_records (across files in order)
    and returns the last header info (ship, city, commander) found before this entry.
    Returns defaults if none found.
    """
    for i in range(index - 1, -1, -1):
        header = extract_header_info(all_records[i]["Notes"])
        if header and header["ship"]:
            return header
    return {"ship": "Unknown/Not Found", "city": "-", "cmd": "-"}


def process_loyallist_comparison(excel_path, word_dir, output_file):
    print(f"Loading {excel_path}...")
    try:
        df_master = pd.read_excel(excel_path)
    except Exception as e:
        print(f"Excel Load Error: {e}")
        return

    # Clean master Excel data
    df_master['Ship_Name'] = df_master['Ship_Name'].astype(str).str.strip().str.lower()
    df_master['Book'] = df_master['Book'].astype(str).str.strip()
    df_master['Notes'] = df_master['Notes'].fillna("").astype(str).str.strip()

    # Build lookup sets
    df_master['Ship_Book_Key'] = df_master['Ship_Name'] + "||" + df_master['Book']
    existing_ship_book_combos = set(df_master['Ship_Book_Key'].unique())
    existing_notes = set(df_master['Notes'].unique())

    print(f"Scraping Word documents in '{word_dir}' (ordered)...")
    all_records = get_word_content_ordered(word_dir)

    missing_records = []
    seen_in_report = set()

    for idx, record in enumerate(all_records):
        note_text = record["Notes"]
        source_file = record["Source_Word_File"]
        mapped_book = FILE_TO_BOOK_MAP.get(source_file, "Unknown Book")

        # Try to extract header info directly from this entry
        extracted_header = extract_header_info(note_text)
        extracted_ship = extracted_header["ship"] if extracted_header else None

        # --- Validation Step 1: Check Ship+Book combo ---
        if extracted_ship:
            ship_book_key = f"{extracted_ship.lower()}||{mapped_book}"
            if ship_book_key in existing_ship_book_combos:
                continue  # Already in master, skip

        # --- Validation Step 2: Check if exact Note text exists ---
        if note_text in existing_notes:
            continue  # Already in master, skip

        # --- Entry is missing from master: determine ship, commander, city ---
        if extracted_ship:
            # Case 1: Entry itself is a ship header line
            final_ship = extracted_ship
            final_commander = extracted_header["cmd"]
            final_city = extracted_header["city"]
            ship_source = "Extracted from entry"
        else:
            # Case 2: Not a ship line — backtrack to find the last header
            last_header = find_last_header_before(idx, all_records)
            final_ship = last_header["ship"]
            final_commander = last_header["cmd"]
            final_city = last_header["city"]
            ship_source = "Backtracked from previous entries"

        report_key = (note_text, source_file, final_ship)
        if report_key not in seen_in_report:
            missing_records.append({
                "Notes": note_text,
                "Source_Word_File": source_file,
                "Ship": final_ship,
                "Commander": final_commander,
                "Arrival_Port_City": final_city,
                "Ship_Source": ship_source
            })
            seen_in_report.add(report_key)

    if missing_records:
        df_final = pd.DataFrame(missing_records)
        df_final = df_final[["Notes", "Source_Word_File", "Ship", "Commander", "Arrival_Port_City", "Ship_Source"]]
        df_final.to_excel(output_file, index=False)
        print(f"\nAnalysis complete. Found {len(missing_records)} records missing from the Master Excel.")
        print(f"Report saved as: {output_file}")
    else:
        print("\nAll entries accounted for. No missing records found based on Ship/Book validation.")


# --- File Paths ---
DIR_NAME = "Book_of_Negroes"
MASTER_FILE = "Black_Loyalist_Directory_Consolidated.xlsx"
REPORT_FILE = "Validated_Missing_Records.xlsx"

if __name__ == "__main__":
    process_loyallist_comparison(MASTER_FILE, DIR_NAME, REPORT_FILE)